"""生成大赛讲解稿 Word 文档"""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import os

doc = Document()

# ========== 页面设置 ==========
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.8)

style = doc.styles['Normal']
style.font.name = '微软雅黑'
style.font.size = Pt(11)
style.paragraph_format.line_spacing = 1.6
style.paragraph_format.space_after = Pt(6)

# ========== 辅助函数 ==========
def add_title(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
    p.paragraph_format.space_after = Pt(4)

def add_subtitle(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(0x63, 0x66, 0xf1)
    p.paragraph_format.space_after = Pt(16)

def add_h1(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
    p.paragraph_format.space_before = Pt(20)
    p.paragraph_format.space_after = Pt(8)

def add_h2(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(13)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x63, 0x66, 0xf1)
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)

def add_body(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(11)
    p.paragraph_format.first_line_indent = Cm(0.7)

def add_bullet(text, indent=0):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    p.paragraph_format.left_indent = Cm(1.0 + indent * 0.8)
    p.paragraph_format.space_after = Pt(2)

def add_highlight_box(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.0)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x63, 0x66, 0xf1)

def add_separator():
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('─' * 40)
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0xcc, 0xcc, 0xcc)
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(10)


# ================================================================
# 封面
# ================================================================
for _ in range(4):
    doc.add_paragraph()

add_title('巴蜀文化双语主持教学智能体')
add_subtitle('—— AI 赋能跨文化传播人才培养的探索与实践')

for _ in range(2):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('四川省教育厅 · 第二届教师人工智能应用能力大赛')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('方向二：AI 教育智能体设计与应用')
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x63, 0x66, 0xf1)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('2026年7月')
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

doc.add_page_break()

# ================================================================
# 时间分配提示
# ================================================================
add_h1('讲解时间分配（10分钟）')
add_bullet('1. 背景与痛点 —— 1.5 分钟')
add_bullet('2. 设计理念与架构 —— 2 分钟')
add_bullet('3. 核心功能展示（七大场景） —— 3.5 分钟')
add_bullet('4. 创新点与实践成效 —— 1.5 分钟')
add_bullet('5. 总结与展望 —— 1.5 分钟')

add_separator()

# ================================================================
# 第一部分：背景与痛点
# ================================================================
add_h1('一、背景与痛点（1.5分钟）')

add_body('各位评委老师好！今天我展示的是一个面向《英语主持与传播》课程的 AI 教育智能体——以巴蜀文化国际传播为核心内容载体，构建"学-练-评-踪"一体化的双语主持教学系统。')

add_h2('教学痛点')

add_body('在双语主持教学中，我们面临三个核心痛点：')

add_bullet('第一，教师精力瓶颈。一个教学班 40 名学生，每人完成一次 3 分钟的双语主持练习，教师需要至少 4 小时才能完成批改——而这只是"内容"部分，还没算上语音、台风、情感表达的指导。')
add_bullet('第二，跨文化教学资源稀缺。巴蜀文化国际传播需要精准的中英双语知识支撑——三星堆怎么讲？川剧变脸怎么翻译？大熊猫保护的国际话语怎么构建？这些高度专业化的教学内容，远非通用教材所能覆盖。')
add_bullet('第三，口语练习缺乏即时反馈。学生课后自主练习主持口语时，完全不知道自己发音对不对、内容好不好，练习变成"黑箱"，效果无从验证。')

add_h2('破题思路')

add_body('我们的答案是：构建一个 AI 教育智能体，让 AI 负责内容逻辑的客观分析——论点是否成立、论据是否充分、文化是否准确、结构是否连贯；而语音的温度、情感的细腻、台上的举手投足——这些需要人文感知的部分，留给最懂学生的老师。这是"人机共生"而非"机器替代"。')

add_separator()

# ================================================================
# 第二部分：设计理念
# ================================================================
add_h1('二、设计理念与架构（2分钟）')

add_h2('2.1 核心理念：AI 评内容，教师评表达')

add_body('这个智能体的定位不是替代教师，而是为教师分担最费时的"内容批改"重复劳动——让教师把精力集中在更高价值的表达指导与人文陪伴上。这与大赛提倡的"人机共生、氛围浸润的教学新生态"高度一致。')

add_h2('2.2 技术架构')

add_body('系统采用前后端分离架构：前端 React + TypeScript SPA，后端 Python FastAPI，知识库采用 ChromaDB 向量存储 + JSON 结构化数据，AI 引擎为 DeepSeek + 讯飞语音评测双模型驱动。')

add_body('核心创新在于 Agent Workflow 设计——系统将评析任务拆解为八个并行维度（文化准确性、观点立场、论据展开、文化转译质量、受众意识、逻辑连贯、叙事结构、跑题检测），每个维度由一个独立的 AI Agent 负责，并行处理后由综合 Agent 汇总判定等级。这大幅提升了评析效率——从教师人工批改的平均 30 分钟/篇，压缩到 30-60 秒/篇。')

add_h2('2.3 知识图谱驱动')

add_body('智能体内置了覆盖巴蜀文化六大模块（自然遗产、历史文脉、非遗技艺、城市精神、当代创新、文化比较）的 25 条双语知识条目，每条包含关键事实、术语表、常见误解和叙事角度建议。评析时通过 RAG 检索相关知识点作为 Ground Truth，确保文化评析的准确性。题库包含 148 道练习题目，覆盖六种主持题型、五个难度等级。')

add_separator()

# ================================================================
# 第三部分：核心功能展示
# ================================================================
add_h1('三、核心功能展示——七个维度融合应用（3.5分钟）')

add_body('下面从大赛要求的"教、学、评、研、管、育、建"七个维度，展示智能体的具体应用场景。')

# --- 教 ---
add_h2('3.1【教】知识图谱驱动的文化教学')

add_body('传统双语主持教学中，巴蜀文化知识依赖教师个人积累。智能体将 25 条结构化双语知识条目向量化存入 ChromaDB，学生在选题时自动匹配相关知识点——包括中英术语对照、关键数据、常见跨文化误解和推荐叙事角度。')

add_highlight_box('应用效果：知识检索响应 < 1 秒，文化准确度评析有了可验证的 Ground Truth。')

# --- 学 ---
add_h2('3.2【学】双轨练习：文字 + 语音')

add_body('智能体提供"文字测评"和"录音测评"两套独立练习链路。文字测评路径：选题 → 知识辅助 → 撰写中英双语主持稿 → 提交八维评析。录音测评路径：从 15 篇巴蜀文化双语文稿中选择一篇 → 录制中英文朗读音频 → 讯飞引擎评测发音准确度、流利度和完整度。两套链路互不干扰，学生可根据学习目标自主选择。')

add_highlight_box('应用效果：实现了"个性化学习路径"——基础薄弱的学生从录音跟读开始，进阶学生挑战即兴串场和观点辩论。')

# --- 评 ---
add_h2('3.3【评】八维深度评析 + 语音评测')

add_body('这是智能体的核心引擎。每次练习完成后，AI 从八个维度对文稿内容进行系统评析，输出 S/A/B/C/D 五档等级和每个维度的详细反馈——包括具体问题定位、修改建议和参考范例。同时，讯飞语音评测引擎给出中文和英文的发音总分、准确度、流利度和完整度四项指标。')

add_highlight_box('应用效果：评析效率提升 30 倍以上，且反馈精准到具体句段——"文化转译质量"维度能识别出学生将"巴蜀包容性"生硬直译为"Ba-Shu inclusiveness"而缺少英文受众所需的解释性展开。')

# --- 研 ---
add_h2('3.4【研】数据驱动的教学研究')

add_body('智能体的成长追踪仪表盘为教学研究提供了定量数据基础：八维能力雷达图追踪每个维度的历史变化、等级分布饼图呈现全班整体水平、成长曲线记录个体进步轨迹、薄弱维度自动预警。教师可以基于这些数据开展"双语主持人能力发展规律""跨文化转译能力的培养路径"等实证研究。')

add_highlight_box('应用效果：从"经验判断"到"数据驱动"——教师可以精准定位全班共性薄弱点（如"受众意识"维度全班均分最低），据此调整教学重点。')

# --- 管 ---
add_h2('3.5【管】教学管理增效')

add_body('教师可以通过仪表盘快速了解全班学情：练习频次统计反映学生投入度、模块覆盖热力图显示教学内容均衡性、等级分布揭示教学效果。可视化数据看板替代了传统的 Excel 手动统计——一次登录，全局掌控。')

# --- 育 ---
add_h2('3.6【育】高阶思维培养')

add_body('智能体的选题设计并非停留在知识记忆层面。148 道题目中，文化解说题（28道）训练文化理解力，即兴串场题（20道）培养临场应变力，跨文化评论题（21道）发展批判性思维，观点辩论题（20道）锻炼论证与反驳能力。从"知道巴蜀文化"到"能用英语讲述巴蜀文化"再到"能在跨文化语境中评论和辩论巴蜀文化"——智能体推动学生实现从识记到创造的高阶思维跃迁。')

add_highlight_box('应用效果：题目的认知层级分布符合布鲁姆教育目标分类学——从记忆、理解到应用、分析、评价、创造。')

# --- 建 ---
add_h2('3.7【建】可持续的知识库建设机制')

add_body('智能体的巴蜀文化知识库采用 JSON 结构化存储，支持教师持续扩充——新增一条知识条目只需编辑 JSON 文件，系统自动向量化索引。题库同样支持教师自定义添加题目。这套"低代码建设"机制确保智能体不是一次性交付的封闭系统，而是可以随着课程发展持续生长的活态资源平台。')

add_separator()

# ================================================================
# 第四部分：创新点
# ================================================================
add_h1('四、创新点与实践成效（1.5分钟）')

add_h2('4.1 三个核心创新')

add_bullet('创新一：学科知识图谱驱动的精准评析。不同于通用写作批改工具，本智能体基于巴蜀文化知识图谱进行文化准确性验证——能识别"变脸是魔术""大熊猫只生活在四川"等常见误解，这是通用 AI 做不到的。')
add_bullet('创新二："AI 评内容 + 教师评表达"的人机协同范式。明确划分 AI 与教师的能力边界——八维评析聚焦内容逻辑，语音情感留给教师——这一分工理念在同类产品中尚属首创。')
add_bullet('创新三：双轨独立练习链路。文字内容测评与录音语音测评各自独立、互不干扰，分别服务于"主持稿写作"和"口语表达"两个不同的教学目标，在同领域智能体中具有差异化优势。')

add_h2('4.2 实践应用数据')

add_body('截至目前，智能体已构建巴蜀文化双语知识库 25 条条目，练习题库 148 道题目（覆盖六大文化模块、六种主持题型、五个难度等级），支持八维内容评析 + 四项语音评测。知识检索响应时间 < 1 秒，并行评析完成时间约 30-60 秒。系统采用 DeepSeek API + 讯飞 ISE WebSocket 流式接口，API 调用成功率 > 98%。')

add_separator()

# ================================================================
# 第五部分：总结与展望
# ================================================================
add_h1('五、总结与展望（1.5分钟）')

add_body('这个智能体的核心价值，不只是"让批改变快了"。更深层的意义在于——')

add_bullet('它让每一个学生在课后都能获得即时、精准、个性化的反馈，不再依赖"有没有老师在"；')
add_bullet('它让巴蜀文化国际传播的教学有了可验证的知识底座，不再依赖教师个人积累的深浅；')
add_bullet('它让人机各自做最擅长的事——AI 做逻辑分析，人做情感判断——构建起"人机共生、氛围浸润"的教学新生态。')

add_body('未来，我们计划在三个方向上持续迭代：第一，将知识库从当前 25 条扩充至 200+ 条目，向量检索升级为混合检索（语义 + 关键词）；第二，增加"多模态主持练习"——上传演讲视频，AI 同时分析内容、语音和肢体语言；第三，沉淀练习数据，训练面向双语主持教学的专用小模型，进一步降低 API 调用成本。')

add_body('感谢各位评委！期待您的指导与建议。')

add_separator()

# ================================================================
# 附录
# ================================================================
add_h1('附录：技术指标一览')

add_bullet('前端框架：React 18 + TypeScript + Vite + Tailwind CSS + Recharts')
add_bullet('后端框架：Python FastAPI + SQLAlchemy + SQLite')
add_bullet('AI 引擎：DeepSeek API（文本评析）+ 讯飞 ISE WebSocket（语音评测）')
add_bullet('知识库：ChromaDB 向量检索 + JSON 结构化种子数据')
add_bullet('巴蜀文化知识条目：25 条（双语句式，含关键事实、术语、误解纠正、叙事角度）')
add_bullet('题库规模：148 道（6 种题型 × 6 大模块 × 5 级难度）')
add_bullet('评析维度：8 维内容（并行 Agent）+ 4 维语音（讯飞 ISE）')
add_bullet('评析速度：并行八维 < 60 秒/篇')
add_bullet('语音评测：中文 + 英文，独立评分')

# ========== 保存 ==========
output_path = os.path.join(os.path.dirname(__file__), '大赛讲解稿_巴蜀文化双语主持教学智能体.docx')
doc.save(output_path)
print(f'文档已生成: {output_path}')
